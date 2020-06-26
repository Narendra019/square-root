# -*- coding: utf-8 -*-
"""
Created on Sun Jun 21 18:11:37 2020

@author: Narendra Chowdary
""
"""

from tkinter import *
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import PIL
import os
import glob
import os.path
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.shared import Pt,Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn

# from docx import paragraph_format.space_before

size= 0
ext = ''
gui = tk.Tk()
gui.geometry("670x240")
gui.configure(bg='mediumspringgreen')
# gui.configure(tk,bg= 'lightgrey')
gui.title("Automate Document")

def getFolderPath():
    folder_selected = filedialog.askopenfile()
    global ext
    try:
      ext = (os.path.basename(folder_selected.name)).rsplit('.', 1)[1]
      folder_selected = os.path.dirname(folder_selected.name)
      folderPath.set(folder_selected)
      os.chdir(folder_selected)
      print(folder_selected)
    except AttributeError:
      pass
     
def doStuff():
    folder = folderPath.get()
    print("Doing stuff with folder",  folder)
    do=TextVal.get()
    os.system('start {0}{1}'.format(do, '.docx'))

    
def StartBtn():
    document = Document()
    section = document.sections[0]
    style = document.styles['Normal']
    font = style.font
    header = section.header
    header.is_linked_to_previous=False
    paragraph = header.paragraphs[0]
    font.size = Pt(35)
    paragraph.text =TextVal.get()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sec_pr = document.sections[0]._sectPr #
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')#single line
        border_el.set(qn('w:sz'), '24') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '26')
        border_el.set(qn('w:color'), 'blue')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) 
    # document.add_paragraph( 'Inspection File is ready',style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    global size1,size2
    real=n.get()
    if (real==2):
      sectPr = section._sectPr
      cols = sectPr.xpath('./w:cols')[0]
      cols.set(qn('w:num'),'1')  
      size1=6.2362204724
      size2=4.0787402
    elif (real==3):
       size1=5.354331
       size2=2.771654
    elif (real==6):
       sectPr = section._sectPr
       cols = sectPr.xpath('./w:cols')[0]
       cols.set(qn('w:num'),'2')
       size1=2.75591
       size2=2.75591
   
    for file in glob.glob('*.%s'%ext):
        img = Image.open(file)
        img= img.resize((1024,768))
        img.save(file)
        paragraph_format = paragraph.paragraph_format
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(file,width=Inches(size1),height=Inches(size2))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        filenam=TextVal.get()
        document.save(filenam+".docx")
   
    print(real)
    messagebox.showinfo("COMPRESS & SAVE ", "YOUR DOCUMENT IS READY")
    print('Done')


def clear():
    TextVal.set(0)
    E.delete(0,END)
    H.delete(0,END)
def delete():
    filenam=TextVal.get()
    try:
      os.remove(filenam+".docx")
      messagebox.showinfo("DETAILS", "DELETED")
    except (FileNotFoundError, PermissionError) :
      messagebox.showerror("Details", "Please Enter the File Tag ")
def close():
    sys.exit()
    
folderPath = StringVar()        
# filepath=folderPath.get()

TextVal=StringVar()

frame0=tk.Frame(gui,bg="lightcyan",highlightthickness=1.0,highlightbackground="black")
frame0.place(relwidth=0.96,relheight=0.95,relx=0.02,rely=0.02)  

E = Entry(frame0,textvariable= folderPath,font=('calbiri',11,'bold'))
E.place(relwidth=0.7,relheight=0.15,relx=0.02,rely=0.06)

l =tk.Label(frame0,text=" CUSTOMER NAME & MR NO :",font=('calbiri',10,'bold'),background='lightcyan')
l.place(relwidth=0.3,relheight=0.15,relx=0.01,rely=0.21)

H = Entry(frame0,textvariable= TextVal,font=('calbiri',11,'bold'))
H.place(relwidth=0.7,relheight=0.13,relx=0.29,rely=0.23)

c =tk.Label(frame0,text="IMAGES PER PAGE",font=('Arial',10,'bold'),background='lightcyan')
c.place(relwidth=0.3,relheight=0.15,relx=0.01,rely=0.39)

n =IntVar() 
cm=ttk.Combobox(frame0,textvariable = n,font=('Arial',12,'bold')) 
   
# Adding combobox drop down list 
cm['value'] =(2,3,6)
cm.set(2)
cm.place(relwidth=0.3,relheight=0.15,relx=0.29,rely=0.4)

btnFind = tk.Button(gui,text="SELECT FOLDER",command=getFolderPath,borderwidth=3,font=('Arial',10,'bold'),background="mediumspringgreen",activebackground="lightblue")
btnFind.place(relwidth=0.20,relheight=0.15,relx=0.73,rely=0.06)

c = tk.Button(frame0 ,text="OPEN", command=doStuff,highlightthickness=5.0,highlightbackground="black",font=('Arial',10,'bold'),background="mediumspringgreen",activebackground="lightblue")
c.place(relwidth=0.20,relheight=0.15,relx=0.03,rely=0.7)

btnStart= tk.Button(frame0, text="RESIZE&DOC ",highlightthickness=1.0,highlightbackground="black",font=('Arial',10,'bold'),background="mediumspringgreen",command=StartBtn,activebackground="lightblue")
btnStart.place(relwidth=0.20,relheight=0.15,relx=0.26,rely=0.7)

btnStart= tk.Button(frame0, text="CLEAR",command=clear,highlightthickness=1.0,highlightbackground="black",font=('Arial',10,'bold'),background="mediumspringgreen",activebackground="lightblue")
btnStart.place(relwidth=0.20,relheight=0.15,relx=0.5,rely=0.7)

btndel= tk.Button(frame0, text="DELETE",command=delete,highlightthickness=1.0,highlightbackground="black",font=('Arial',10,'bold'),background="mediumspringgreen",activebackground="lightblue")
btndel.place(relwidth=0.20,relheight=0.15,relx=0.75,rely=0.7)

gui.resizable(0,0)


gui.mainloop()




